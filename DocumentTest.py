from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn

class DocumentTest:

    def __init__(self):
        self.document = Document()

    def add_para(self, text, font='宋体', size=14, before=None, after=None, isBold=False, isCenter=False):
        p = self.document.add_paragraph(text)
        run = p.runs[0]
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
        run.font.size = Pt(size)
        if isBold:
            run.bold = True
        if isCenter:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if before is not None:
            p.paragraph_format.space_before = Pt(before)
        if after is not None:
            p.paragraph_format.space_after = Pt(after)
        return p

    def add_text(self, text, font='宋体', size=14, isBold=False, isCenter=False, paragraph=None):
        if paragraph is None:
            paragraph = self.document.add_paragraph()
        run = paragraph.add_run(text)
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
        run.font.size = Pt(size)
        if isBold:
            run.bold = True
        if isCenter:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return paragraph

    
    def save(self, path):
        self.document.save(path)

if __name__ == "__main__":
    docPlayer = DocumentTest()
    docPlayer.add_para('\n\n\n后疫情时代网络热词和社会现象', size=16, font="黑体", isBold=True, isCenter=True)
    docPlayer.add_para("\n张三", size=14, isCenter=True)
    docPlayer.add_para("(中南海)", size=12, isCenter=True)
    docPlayer.add_text("本文探讨了网络热词的分类和变迁...", size=10.5, font="仿宋",paragraph=docPlayer.add_para("\t摘要", size=10.5, font="仿宋", isBold=True))
    docPlayer.save("example.docx")
